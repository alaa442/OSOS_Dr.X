import os
from docx import Document
import fitz  # from PyMuPDF
import pandas as pd
import openpyxl


def extract_docx(file_path):
    doc = Document(file_path)
    text = []

    for para in doc.paragraphs:
        if para.text.strip():
            text.append(para.text.strip())

    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_text:
                text.append(" | ".join(row_text))

    return "\n".join(text)


def extract_pdf(file_path):
    doc = fitz.open(file_path)
    pages = []

    for i, page in enumerate(doc):
        text = page.get_text("text")
        pages.append(f"--- Page {i+1} ---\n{text}")

    return "\n".join(pages)



def extract_all_excel_data(file_path):
    """
    Extracts all readable text from all sheets in an Excel file (.xlsx, .xlsm),
    including sub-tables and cells with formulas. Hidden sheets will also be read.
    """
    wb = openpyxl.load_workbook(file_path, data_only=True)
    extracted = []

    for sheet in wb.worksheets:
        extracted.append(f"\n--- Sheet: {sheet.title} ---\n")
        current_table = []

        for row in sheet.iter_rows(values_only=True):
            if row is None or all(cell is None for cell in row):
                if current_table:
                    extracted.append("\n".join(current_table))
                    extracted.append("\n--- Subtable Break ---\n")
                    current_table = []
            else:
                row_values = [str(cell).strip() for cell in row if cell is not None]
                if row_values:
                    current_table.append(" | ".join(row_values))

        # Add any remaining table after loop
        if current_table:
            extracted.append("\n".join(current_table))

    return "\n".join(extracted)


def extract_text(file_path):
    ext = os.path.splitext(file_path)[1].lower()

    if ext == ".docx":
        return extract_docx(file_path)
    elif ext == ".pdf":
        return extract_pdf(file_path)
    elif ext in [".xlsx", ".xls", ".xlsm"]:
        return extract_all_excel_data(file_path)
    else:
        raise ValueError(f"Unsupported file type: {ext}")
